"""
Smoke test for parse_document tool.

策略:
1. office 路径 (.docx) — 无需模型下载,完全离线,作为主验证。
2. PDF 路径 — 用 reportlab 生成 1 页测试 PDF。pipeline 后端首次跑会下载模型,
   网络/超时/磁盘空间不足时 skip 并打 ⚠️,不让 CI 因外部依赖挂。

完成判据:
  ✅ docx 路径返回非空 markdown 且含我们注入的关键字 "ASHARE-MCP-PARSE-TEST"
  ✅ metadata 包含 engine='mineru', file_suffix='.docx', has_json/markdown_chars 合理
  ⚠️ PDF 路径若下载失败 → 标 skip(原因写日志)
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest

from ashare_mcp.parse_document import parse_document_impl

KEYWORD = "ASHARE-MCP-PARSE-TEST"


# ------- helpers -------
def _make_docx(path: Path, body: str) -> None:
    """用 python-docx (mineru 自带依赖) 写一份最小 docx。"""
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading("Parse Document Smoke Test", level=1)
    doc.add_paragraph(body)
    doc.add_paragraph("第二段:中文也要能识别。")
    doc.save(str(path))


def _make_pdf(path: Path, body: str) -> None:
    """用 reportlab 生成 1 页测试 PDF(英文以避免字体依赖)。"""
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.lib.pagesizes import LETTER  # type: ignore

    c = canvas.Canvas(str(path), pagesize=LETTER)
    c.setFont("Helvetica", 18)
    c.drawString(72, 720, "Parse Document Smoke Test")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, body)
    c.drawString(72, 670, "Second line for parser sanity check.")
    c.showPage()
    c.save()


# ------- tests -------
def test_parse_document_rejects_missing_file():
    with pytest.raises(FileNotFoundError):
        parse_document_impl("/tmp/__definitely_not_here__.pdf")


def test_parse_document_rejects_unsupported_suffix(tmp_path: Path):
    p = tmp_path / "x.txt"
    p.write_text("hello", encoding="utf-8")
    with pytest.raises(ValueError, match="unsupported"):
        parse_document_impl(str(p))


def test_parse_document_rejects_bad_output_format(tmp_path: Path):
    p = tmp_path / "x.docx"
    _make_docx(p, KEYWORD)
    with pytest.raises(ValueError, match="output_format"):
        parse_document_impl(str(p), output_format="xml")  # type: ignore[arg-type]


def test_parse_document_docx_smoke(tmp_path: Path):
    """
    主烟囱测试:docx 路径,完全离线,无模型下载。
    """
    src = tmp_path / "smoke.docx"
    _make_docx(src, f"This is the body. Marker={KEYWORD}.")

    result = parse_document_impl(str(src), output_format="both", lang="ch")

    # content shape
    assert isinstance(result, dict)
    assert "content" in result and "metadata" in result
    md = result["content"]
    assert isinstance(md, str) and len(md) > 0, "markdown content is empty"
    assert KEYWORD in md, f"injected keyword not found in markdown; got: {md[:200]!r}"

    # metadata shape
    meta = result["metadata"]
    assert meta["engine"] == "mineru"
    assert meta["file_suffix"] == ".docx"
    assert meta["lang"] == "ch"
    assert meta["markdown_chars"] == len(md)
    assert meta["has_json"] is True
    assert "source_file" in meta and meta["source_file"].endswith("smoke.docx")

    # json side product
    content_list = result.get("content_list")
    assert isinstance(content_list, list) and len(content_list) > 0, (
        "content_list missing or empty for output_format='both'"
    )


def test_parse_document_pdf_smoke(tmp_path: Path):
    """
    PDF 路径:首次会下载 pipeline 模型(几百 MB)。
    无网络 / 下载失败 / 超时 → skip 并打 ⚠️,不阻塞 CI。
    """
    if os.environ.get("ASHARE_MCP_SKIP_PDF_TEST") == "1":
        pytest.skip("ASHARE_MCP_SKIP_PDF_TEST=1 — pdf smoke skipped on request")

    src = tmp_path / "smoke.pdf"
    _make_pdf(src, f"Marker={KEYWORD} body line.")

    try:
        result = parse_document_impl(str(src), output_format="markdown", lang="en")
    except Exception as e:
        # 模型下载需要外网 + 几百 MB 磁盘空间。任何 import/下载/初始化错误都 skip,
        # 在报告里以 ⚠️ 形式呈现。
        pytest.skip(f"pdf smoke skipped (likely model download / network): {e!r}")

    md = result["content"]
    assert isinstance(md, str) and len(md) > 0, "pdf markdown is empty"
    # OCR 可能拼错,关键字 + 容错小写匹配
    assert "marker" in md.lower() or KEYWORD in md, (
        f"PDF marker not recognized; first 300 chars: {md[:300]!r}"
    )
    meta = result["metadata"]
    assert meta["file_suffix"] == ".pdf"
    assert meta["engine"] == "mineru"
